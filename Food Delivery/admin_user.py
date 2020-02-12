# Author: Natnael Alemayehu / Mahalinoro Razafimanjato

# Importing the class files needed
import docx

from user_class import User
import food_class
import drink_class
import sys
import main_application

# Dictionary that will contain the admin users after creation
admin_user = {}


# This is the class admin user with the instance variables and the __init__ method
class AdminUser(User):
    def __init__(self, name):
        super().__init__(name)


def view_items():
    # This function will display the food items and the drinks item to the user
    print("You have this foods: ")
    for food in food_class.food_item.keys():
        print(food)
    print("You have this drinks: ")
    for drink in drink_class.drink_item.keys():
        print(drink)
    admin_menu()


# This function will add new food item in the storage
def add_food_item(food_name, food_price, food_size, food_description):
    if food_name == "" or food_price == "" or food_size == "" or food_description == "":
        raise ValueError("Empty value")
    food_class.food_item.update({food_name: food_class.Food(food_name,
                                                            food_price, food_size, food_description)})
    print("Successfully added food item")
    admin_menu()


# This function will add new drink in the storage
def add_drink_item(drink_name, drink_price, drink_volume, drink_stock, drink_description):
    if drink_name == "" or drink_price == "" or drink_volume == "" or drink_stock == "" or drink_description == "":
        raise ValueError("Empty value")
    drink_class.drink_item.update({drink_name: drink_class.Drink(
        drink_name, drink_price, drink_volume, drink_stock, drink_description)})
    print("Successfully added food item")
    admin_menu()


# This function will remove a selected food item
def remove_food_item(food_to_delete):
    if food_to_delete == "":
        raise ValueError("Empty value")
    if food_to_delete in food_class.food_item.keys():
        food_class.food_item.pop(food_to_delete)
        print("Item removed successfully")
        admin_menu()


# This function will remove a selected drink item
def remove_drink_item(drink_to_delete):
    if drink_to_delete == "":
        raise ValueError("Empty value")
    if drink_to_delete in drink_class.drink_item.keys():
        drink_class.drink_item.pop(drink_to_delete)
        print("Item removed successfully")
        admin_menu()


# This function will update the information of a selected food item
def update_food_information(food_to_update, update_choice):
    if food_to_update == "" or update_choice == "":
        raise ValueError("Empty value")
    if update_choice == "name":
        updated_name = input(
            "What do you want to rename it as? \n")
        food_class.food_item[food_to_update].name = updated_name
    elif update_choice == "price":
        updated_price = input("What do you want to update it as? \n")
        try:
            updated_price = int(updated_price)
        except ValueError:
            print("Only strings allowed")
            admin_menu()
        food_class.food_item[food_to_update].price = updated_price
    elif update_choice == "size":
        updated_size = input("What do you want to update it as? \n")
        food_class.food_item[food_to_update].size = updated_size
    elif update_choice == "description":
        updated_description = input(
            "What do you want to update it as? \n")
        food_class.food_item[food_to_update].description = updated_description
    print("Item updated successfully")
    admin_menu()


# This function will update the information of a selected drink item
def update_drink_information(drink_to_update, update_choice):
    if drink_to_update == "" or update_choice == "":
        raise ValueError("Empty value")
    if update_choice == "name":
        updated_name = input(
            "What do you want to rename it as? \n")
        drink_class.drink_item[drink_to_update].name = updated_name
    elif update_choice == "price":
        updated_price = input("What do you want to update it as? \n")
        try:
            updated_price = int(updated_price)
        except ValueError:
            print("Only strings allowed")
            admin_menu()
        drink_class.drink_item[drink_to_update].price = updated_price
    elif update_choice == "volume":
        updated_volume = input("What do you want to update it as? \n")
        try:
            updated_volume = int(updated_volume)
        except ValueError:
            print("Only strings allowed")
            admin_menu()
        drink_class.drink_item[drink_to_update].volume = updated_volume
    elif update_choice == "description":
        updated_description = input(
            "What do you want to update it as? \n")
        drink_class.drink_item[drink_to_update].description = updated_description
    print("Item updated successfully")
    admin_menu()


# This function will generate the revenue made from drink and food sold so far
def revenue_generated():
    food_revenue = 0
    for food in food_class.food_item.keys():
        food_revenue += food_class.food_item[food].price * food_class.food_item[food].item_count_delivered
    print("Total revenue from food is", food_revenue)
    drink_revenue = 0
    for drink in drink_class.drink_item.keys():
        drink_revenue += drink_class.drink_item[drink].price * drink_class.drink_item[drink].item_count_delivered
    print("Total revenue from drink is", drink_revenue)
    net_revenue = 0
    net_revenue = food_revenue + drink_revenue
    print("Total revenue is", net_revenue)
    admin_menu()


# This function will display the number of order for food and drinks item
def order_number():
    user_input = input('''Press 1 to view order number for food 
    Press 2 to view order number for drink \n''')
    food_item_count = 0
    drink_item_count = 0
    if user_input == "1":
        for food in food_class.food_item.keys():
            print("{} orders received for {}".format(str(food_class.food_item[food].item_count_delivered),
                                                     food_class.food_item[food].name))
        admin_menu()
    elif user_input == "2":
        pass
        for drink in drink_class.drink_item.keys():
            print("{} orders received for {}".format(str(drink_class.drink_item[drink].item_count_delivered),
                                                     drink_class.drink_item[drink].name))
        admin_menu()
    elif user_input == "3":
        for food in food_class.food_item.keys():
            food_item_count += food_class.food_item[food].item_count_delivered
        for drink in drink_class.drink_item.keys():
            drink_item_count += drink_class.drink_item[drink].item_count_delivered

        print("{} orders for both food and drink".format(str(food_item_count + drink_item_count)))
        admin_menu()


# This function is the main menu for admin user
# The admin menu will let the user view item, add new item, remove item, update information, see revenue and see
# order numbers as well as they can log out of the system if they wish

def admin_menu():
    document = docx.Document('user_session.docx')
    print('----------- Welcome to E-MEAL MENU ---------')
    print("Press 1 to view your item in the storage ")
    print("Press 2 to add new item in the storage ")
    print("Press 3 to remove item in the storage")
    print("Press 4 to update information about a specific item")
    print("Press 5 to see revenue generated")
    print("Press 6 to see order numbers for items")
    print("Press L to logout of your account")
    print("Press q or Q to quit.")
    user_choice = input("Type in the corresponding number for you want to do \n")
    if user_choice == "1":
        document.add_paragraph('User selected 1 to view items in the storage.')
        view_items()
    elif user_choice == "2":
        document.add_paragraph('User selected 2 to add new item in the storage.')
        print("What do you want to add?")
        user_input = input("Press 1 to add food and 2 to add drink \n")
        if user_input == "1":
            document.add_paragraph('User pressed 1 to add new food.')
            food_name = input("Food name \n")
            document.add_paragraph('User inserted {} as food name.'.format(food_name))
            food_price = input("price of food \n")
            document.add_paragraph('User inserted {} as food price.'.format(food_price))
            try:
                food_price = int(food_price)
            except ValueError:
                document.add_paragraph('User got errors for inserting not integer price for food.')
                print("Integer only")
                admin_menu()
            food_size = input("Size \n")
            document.add_paragraph('User inserted {} as food size.'.format(food_size))
            food_description = input("description \n")
            document.add_paragraph('User inserted {} as food description.'.format(food_description))
            add_food_item(food_name, food_price, food_size, food_description)
            document.add_paragraph('User added {} successfully in the Food Item.'.format(food_name))
        elif user_input == "2":
            drink_name = input("Drink name \n")
            drink_price = input("price of drink \n")
            drink_volume = input("Drink volume \n")
            drink_stock = input("Stock \n")
            try:
                drink_price = int(drink_price)
                drink_volume = int(drink_volume)
                drink_stock = int(drink_stock)
            except ValueError:
                print("integer only")
                admin_menu()
            drink_description = input("Drink description \n")
            add_drink_item(drink_name, drink_price, drink_volume,
                           drink_stock, drink_description)
        else:
            print("Error try again")
            admin_menu()
        document.save('user_session.docx')
    elif user_choice == "3":
        print("What do you want to delete")
        user_input = input("press 1 to delete food items or 2 for drink items \n")
        if user_input == "1":
            for item in food_class.food_item.keys():
                print(item)
            food_to_delete = input("type in name of food to delete \n")
            remove_food_item(food_to_delete)
        elif user_input == "2":
            for item in drink_class.drink_item.keys():
                print(item)
            drink_to_delete = input("type in name of food to delete \n")
            remove_drink_item(drink_to_delete)
        else:
            print("Error try again")
            admin_menu()
    elif user_choice == "4":
        print("What do you want to update")
        user_input = input(
            "press 1 to update food items or 2 to update drink items \n")
        if user_input == "1":
            for item in food_class.food_item.keys():
                print(item)
            food_to_update = input("type in name of food to update \n")
            if food_to_update in food_class.food_item.keys():
                update_choice = input(
                    "What do you want to update? name, price, size? \n")
                update_food_information(food_to_update, update_choice)
        elif user_input == "2":
            for item in drink_class.drink_item.keys():
                print(item)
            drink_to_update = input("type in name of food to delete \n")
            if drink_to_update in drink_class.drink_item.keys():
                update_choice = input(
                    "What do you want to update? \n name, price, volume? \n")
                update_drink_information(drink_to_update, update_choice)
        else:
            print("Error try again")
            admin_menu()
    elif user_choice == "5":
        revenue_generated()
    elif user_choice == "6":
        order_number()
    elif user_choice == "L":
        main_application.menu()
    elif user_choice == "q" or "Q":
        sys.exit()
    else:
        print("Error, try again!")
        admin_menu()
    document.save('user_session.docx')


if __name__ == '__main__':
    admin_menu()



